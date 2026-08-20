import pytest
from docx import Document

from backend.app.brief.extraction import extract_text
from backend.app.domain.errors import DomainError


def test_extracts_utf8_bom_and_normalizes(tmp_path):
    path = tmp_path / "brief.md"
    path.write_text("\ufeff品牌：清爽\r\n\r\n\r\n场景：通勤\x00", encoding="utf-8")
    assert extract_text(path) == "品牌：清爽\n\n场景：通勤"


def test_extracts_docx_paragraphs_and_tables(tmp_path):
    path = tmp_path / "brief.docx"
    document = Document()
    document.add_paragraph("项目：气泡水")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "品牌"
    table.cell(0, 1).text = "清爽"
    document.save(path)
    assert extract_text(path) == "项目：气泡水\n品牌\t清爽"


def test_rejects_pdf(tmp_path):
    path = tmp_path / "brief.pdf"
    path.write_bytes(b"%PDF")
    with pytest.raises(DomainError) as error:
        extract_text(path)
    assert error.value.status_code == 415
